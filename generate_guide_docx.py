"""Generate the Complete Watcher Fix Guide as a Word (.docx) document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_cell_shading(cell, color):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_paragraph(doc, text, style_name=None, bold=False, italic=False,
                         font_size=None, color=None, alignment=None, space_after=None,
                         space_before=None):
    """Add a paragraph with custom styling."""
    p = doc.add_paragraph()
    if style_name:
        p.style = doc.styles[style_name]
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p


def add_code_block(doc, code_text):
    """Add a code block with gray background."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'F0F0F0')
    p = cell.paragraphs[0]
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)
    doc.add_paragraph()  # spacing after


def add_info_box(doc, lines):
    """Add a blue info box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'E6F3FF')
    # Add blue left border
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:left w:val="single" w:sz="12" w:color="0078C8"/>'
        '</w:tcBorders>'
    )
    tcPr.append(borders)
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 60, 120)
    doc.add_paragraph()


def add_warning_box(doc, text):
    """Add an orange warning box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'FFF5E6')
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:left w:val="single" w:sz="12" w:color="FFA500"/>'
        '  <w:top w:val="single" w:sz="4" w:color="FFA500"/>'
        '  <w:bottom w:val="single" w:sz="4" w:color="FFA500"/>'
        '  <w:right w:val="single" w:sz="4" w:color="FFA500"/>'
        '</w:tcBorders>'
    )
    tcPr.append(borders)
    p = cell.paragraphs[0]
    run = p.add_run("WARNING: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(180, 100, 0)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(180, 100, 0)
    doc.add_paragraph()


def add_step(doc, step_text):
    """Add a green step header."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(step_text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 100, 50)


def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(text)
        run.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)


def create_docx():
    doc = Document()

    # ==================== Set default font ====================
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ==================== COVER PAGE ====================
    for _ in range(4):
        doc.add_paragraph()

    add_styled_paragraph(doc, 'Complete Fix Guide',
                         bold=True, font_size=32, color=(0, 80, 160),
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    add_styled_paragraph(doc, 'All 3 Watchers Authentication',
                         bold=True, font_size=24, color=(40, 40, 40),
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Divider
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('_' * 50)
    run.font.color.rgb = RGBColor(0, 120, 200)
    run.font.size = Pt(12)

    add_styled_paragraph(doc, 'Gmail  |  WhatsApp  |  LinkedIn',
                         font_size=14, color=(80, 80, 80),
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12)

    add_styled_paragraph(doc, 'Personal AI Employee Hackathon 0 - Silver Tier',
                         font_size=12, color=(80, 80, 80),
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

    add_styled_paragraph(doc, 'AI Employee Vault Project',
                         italic=True, font_size=11, color=(120, 120, 120),
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_styled_paragraph(doc, 'February 2026',
                         italic=True, font_size=11, color=(120, 120, 120),
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ==================== ROOT CAUSE ====================
    add_styled_paragraph(doc, 'Root Cause Analysis',
                         bold=True, font_size=20, color=(0, 80, 160),
                         space_after=8)

    p = doc.add_paragraph()
    run = p.add_run(
        'Your watchers cannot sign in because authentication was never set up. '
        'Each service (Gmail, WhatsApp, LinkedIn) requires a one-time OAuth/token '
        'setup before the watcher scripts can connect to their APIs.'
    )
    run.font.size = Pt(11)

    add_styled_paragraph(doc, 'What Was Found',
                         bold=True, font_size=14, color=(40, 40, 40),
                         space_before=12, space_after=6)

    add_bullet(doc, 'Gmail Watcher: ', 'PROBLEM: ')
    p = doc.add_paragraph(
        'The code tries to load gmail_token.json but this file was never created. '
        'The .env still has placeholder value: GOOGLE_REFRESH_TOKEN=your_google_refresh_token'
    )
    p.paragraph_format.left_indent = Cm(1.27)
    for run in p.runs:
        run.font.size = Pt(10)

    add_bullet(doc, 'WhatsApp Watcher: ', 'PROBLEM: ')
    p = doc.add_paragraph(
        'The file whatsapp_watcher.py did NOT EXIST in the watchers/ folder. '
        'Only an MCP server existed with simulated (fake) data.'
    )
    p.paragraph_format.left_indent = Cm(1.27)

    add_bullet(doc, 'LinkedIn Watcher: ', 'PROBLEM: ')
    p = doc.add_paragraph(
        'The file linkedin_watcher.py did NOT EXIST in the watchers/ folder. '
        'Only an MCP server existed with simulated (fake) data. '
        'The .env had all placeholder values.'
    )
    p.paragraph_format.left_indent = Cm(1.27)

    doc.add_paragraph()
    add_warning_box(doc,
        'Your .env file contains what appear to be real API keys (OpenAI, Google, DigitalOcean). '
        'If this repo is public on GitHub, rotate ALL keys immediately!'
    )

    doc.add_page_break()

    # ==================== GMAIL ====================
    add_styled_paragraph(doc, 'WATCHER 1: Gmail - Fix Steps',
                         bold=True, font_size=20, color=(0, 80, 160),
                         space_after=8)

    add_styled_paragraph(doc, 'Problem',
                         bold=True, font_size=14, color=(40, 40, 40), space_after=4)
    p = doc.add_paragraph(
        'No gmail_token.json exists. The Gmail watcher needs an OAuth2 token to authenticate '
        'with Google. The code calls Credentials.from_authorized_user_file() but the file '
        'was never created through the OAuth flow.'
    )

    add_step(doc, 'Step 1: Install Required Packages')
    p = doc.add_paragraph('Open your terminal and run:')
    add_code_block(doc, 'pip install google-auth google-auth-oauthlib google-api-python-client python-dotenv')

    add_step(doc, 'Step 2: Set Up Google Cloud Project')
    p = doc.add_paragraph('Go to https://console.cloud.google.com/ and follow these sub-steps:')
    add_bullet(doc, 'Create a new project (or select an existing one)')
    add_bullet(doc, 'Navigate to: APIs & Services > Library')
    add_bullet(doc, 'Search for "Gmail API" and click Enable')
    add_bullet(doc, 'Go to: APIs & Services > Credentials')
    add_bullet(doc, 'Click "Create Credentials" > select "OAuth Client ID"')
    add_bullet(doc, 'Application type: select "Desktop App"')
    add_bullet(doc, 'Click "Create", then click "Download JSON"')
    add_bullet(doc, 'Save the downloaded file as:')
    add_code_block(doc, r'D:\syeda Gulzar Bano\AI_Employee_Vault_\credentials\client_secret.json')

    add_step(doc, 'Step 3: Configure OAuth Consent Screen')
    p = doc.add_paragraph('Still in Google Cloud Console:')
    add_bullet(doc, 'Go to "OAuth consent screen" in the left sidebar')
    add_bullet(doc, 'User Type: Select "External" (or "Internal" if using Google Workspace)')
    add_bullet(doc, 'Fill in the required fields (App name, user support email)')
    add_bullet(doc, 'Under "Test users", click "Add Users"')
    add_bullet(doc, 'Add YOUR Gmail address as a test user')
    add_bullet(doc, 'Save and continue through all steps')

    add_step(doc, 'Step 4: Run the Auth Setup (One Time Only)')
    p = doc.add_paragraph('This script opens your browser for Google sign-in and creates gmail_token.json:')
    add_code_block(doc,
        r'cd "D:\syeda Gulzar Bano\AI_Employee_Vault_\credentials"' + '\n'
        r'python gmail_auth_setup.py'
    )
    add_info_box(doc, [
        'What happens:',
        '1. A browser window opens automatically',
        '2. Sign in with your Google account',
        '3. Click "Allow" to grant Gmail read access',
        '4. gmail_token.json is created in the credentials/ folder',
        '5. You will see: "SETUP COMPLETE!" in the terminal'
    ])

    add_step(doc, 'Step 5: Run the Gmail Watcher')
    add_code_block(doc,
        r'cd "D:\syeda Gulzar Bano\AI_Employee_Vault_\watchers"' + '\n'
        r'python gmail_watcher.py'
    )
    p = doc.add_paragraph()
    run = p.add_run('You should see: "Gmail API service initialized successfully!" -- This means it is working.')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 130, 0)

    doc.add_page_break()

    # ==================== WHATSAPP ====================
    add_styled_paragraph(doc, 'WATCHER 2: WhatsApp - Fix Steps',
                         bold=True, font_size=20, color=(0, 80, 160),
                         space_after=8)

    add_styled_paragraph(doc, 'Problem',
                         bold=True, font_size=14, color=(40, 40, 40), space_after=4)
    p = doc.add_paragraph(
        'There was NO whatsapp_watcher.py file in the watchers/ folder. A new one was '
        'created that uses the Meta WhatsApp Business Cloud API (the official way to '
        'programmatically access WhatsApp).'
    )

    add_step(doc, 'Step 1: Create a Meta Developer App')
    p = doc.add_paragraph('Go to https://developers.facebook.com/ and:')
    add_bullet(doc, 'Click "My Apps" in the top right, then "Create App"')
    add_bullet(doc, 'Select app type: "Business"')
    add_bullet(doc, 'Fill in app name (e.g., "AI Employee WhatsApp")')
    add_bullet(doc, 'After the app is created, find "WhatsApp" in the products list')
    add_bullet(doc, 'Click "Set Up" next to WhatsApp to add it to your app')

    add_step(doc, 'Step 2: Get Your API Credentials')
    p = doc.add_paragraph('In your app dashboard, navigate to WhatsApp > API Setup:')
    add_bullet(doc,
        'You will see a token at the top. Click "Copy". This is your WHATSAPP_ACCESS_TOKEN.',
        bold_prefix='Temporary Access Token: ')
    add_bullet(doc,
        'Shown below the token. This is your WHATSAPP_PHONE_NUMBER_ID.',
        bold_prefix='Phone Number ID: ')
    add_bullet(doc,
        'Shown in the sidebar or settings. This is your WHATSAPP_BUSINESS_ACCOUNT_ID.',
        bold_prefix='Business Account ID: ')

    add_step(doc, 'Step 3: Update Your .env File')
    p = doc.add_paragraph('Open your .env file and replace the placeholder values with real ones:')
    add_code_block(doc,
        'WHATSAPP_ACCESS_TOKEN=EAAGm0...your_real_token_here\n'
        'WHATSAPP_PHONE_NUMBER_ID=1234567890\n'
        'WHATSAPP_BUSINESS_ACCOUNT_ID=9876543210'
    )

    add_step(doc, 'Step 4: Install Dependencies and Run')
    add_code_block(doc,
        'pip install requests python-dotenv\n'
        r'cd "D:\syeda Gulzar Bano\AI_Employee_Vault_\watchers"' + '\n'
        'python whatsapp_watcher.py'
    )
    p = doc.add_paragraph()
    run = p.add_run('You should see: "WhatsApp credentials validated successfully!"')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 130, 0)

    doc.add_paragraph()
    add_warning_box(doc,
        'Temporary access tokens expire in 24 hours. For permanent use, '
        'submit your Meta app for review and get a permanent token.'
    )

    add_styled_paragraph(doc, 'Optional: Webhook Mode (Real-Time)',
                         bold=True, font_size=13, color=(40, 40, 40),
                         space_before=12, space_after=4)
    p = doc.add_paragraph(
        'For real-time message receiving (instead of polling), you can use webhook mode. '
        'This requires exposing a local server using ngrok:'
    )
    add_code_block(doc,
        '# Install ngrok: https://ngrok.com/download\n'
        'ngrok http 8080\n'
        '\n'
        '# Set the ngrok URL as your webhook URL in Meta Developer Console:\n'
        '# WhatsApp > Configuration > Webhook URL: https://your-url.ngrok.io\n'
        '# Verify Token: ai_employee_verify\n'
        '\n'
        '# Then run watcher in webhook mode:\n'
        'set WHATSAPP_WATCHER_MODE=webhook\n'
        'python whatsapp_watcher.py'
    )

    doc.add_page_break()

    # ==================== LINKEDIN ====================
    add_styled_paragraph(doc, 'WATCHER 3: LinkedIn - Fix Steps',
                         bold=True, font_size=20, color=(0, 80, 160),
                         space_after=8)

    add_styled_paragraph(doc, 'Problem',
                         bold=True, font_size=14, color=(40, 40, 40), space_after=4)
    p = doc.add_paragraph(
        'There was NO linkedin_watcher.py file in the watchers/ folder. A new one was '
        'created that uses the LinkedIn API v2 with OAuth 2.0 authentication.'
    )

    add_step(doc, 'Step 1: Create a LinkedIn Developer App')
    p = doc.add_paragraph('Go to https://www.linkedin.com/developers/ and:')
    add_bullet(doc, 'Click "Create App"')
    add_bullet(doc, 'Fill in: App name, LinkedIn Page (create one if needed), App logo')
    add_bullet(doc, 'Accept the terms and click "Create App"')
    add_bullet(doc, 'Under the "Products" tab, request access to:')
    p = doc.add_paragraph('       - "Sign In with LinkedIn using OpenID Connect"')
    p = doc.add_paragraph('       - "Share on LinkedIn" (for posting capabilities)')

    add_step(doc, 'Step 2: Configure OAuth Settings')
    p = doc.add_paragraph('Under the "Auth" tab of your LinkedIn app:')
    add_bullet(doc, 'Copy the "Client ID" value')
    add_bullet(doc, 'Copy the "Client Secret" value (click the eye icon to reveal)')
    add_bullet(doc, 'Under "OAuth 2.0 settings", click "Add redirect URL"')
    add_bullet(doc, 'Add this exact URL: http://localhost:9090/callback')
    add_bullet(doc, 'Click "Update" to save')

    add_step(doc, 'Step 3: Update .env File')
    p = doc.add_paragraph('Add the Client ID and Secret to your .env:')
    add_code_block(doc,
        'LINKEDIN_CLIENT_ID=86abc1234your_real_id\n'
        'LINKEDIN_CLIENT_SECRET=WPL...your_real_secret'
    )

    add_step(doc, 'Step 4: Run LinkedIn Auth Setup (One Time)')
    add_code_block(doc,
        'pip install requests python-dotenv\n'
        r'cd "D:\syeda Gulzar Bano\AI_Employee_Vault_\credentials"' + '\n'
        'python linkedin_auth_setup.py'
    )
    add_info_box(doc, [
        'What happens:',
        '1. A browser window opens with LinkedIn login',
        '2. Sign in with your LinkedIn account',
        '3. Click "Allow" to grant permissions',
        '4. Browser shows "Success!"',
        '5. Terminal shows your access token',
        '6. Token is saved to credentials/linkedin_token.json'
    ])

    add_step(doc, 'Step 5: Copy Token to .env')
    p = doc.add_paragraph('The auth script prints your token. Copy it to .env:')
    add_code_block(doc, 'LINKEDIN_ACCESS_TOKEN=AQV...the_long_token_from_step4')

    add_step(doc, 'Step 6: Run the LinkedIn Watcher')
    add_code_block(doc,
        r'cd "D:\syeda Gulzar Bano\AI_Employee_Vault_\watchers"' + '\n'
        'python linkedin_watcher.py'
    )
    p = doc.add_paragraph()
    run = p.add_run('You should see: "Authenticated as: Your Name" -- This means it is working.')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 130, 0)

    doc.add_paragraph()
    add_warning_box(doc,
        'LinkedIn access tokens expire in 60 days. Re-run linkedin_auth_setup.py to renew.'
    )

    doc.add_page_break()

    # ==================== SUMMARY ====================
    add_styled_paragraph(doc, 'Summary',
                         bold=True, font_size=20, color=(0, 80, 160),
                         space_after=8)

    add_styled_paragraph(doc, 'Problem & Fix Overview',
                         bold=True, font_size=14, color=(40, 40, 40), space_after=6)

    # Summary table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    headers = ['Watcher', 'Problem', 'Fix']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, '0050A0')
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)

    # Data rows
    data = [
        ['Gmail', 'No OAuth token file', 'Run gmail_auth_setup.py after Google Cloud setup'],
        ['WhatsApp', 'File did not exist + no API token', 'Created whatsapp_watcher.py + get Meta token'],
        ['LinkedIn', 'File did not exist + no API token', 'Created linkedin_watcher.py + run auth setup'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            if row_idx % 2 == 0:
                set_cell_shading(cell, 'F8F8F8')
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(10)
            if col_idx == 0:
                run.bold = True

    doc.add_paragraph()

    # Files table
    add_styled_paragraph(doc, 'Files Created / Fixed',
                         bold=True, font_size=14, color=(40, 40, 40), space_after=6)

    table2 = doc.add_table(rows=6, cols=2)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header
    for i, header in enumerate(['File Path', 'Status']):
        cell = table2.rows[0].cells[i]
        set_cell_shading(cell, '0050A0')
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)

    files = [
        ('credentials/gmail_auth_setup.py', 'NEW - One-time Gmail auth'),
        ('credentials/linkedin_auth_setup.py', 'NEW - One-time LinkedIn auth'),
        ('watchers/gmail_watcher.py', 'FIXED - Proper token + auto-refresh'),
        ('watchers/whatsapp_watcher.py', 'NEW - Meta Cloud API watcher'),
        ('watchers/linkedin_watcher.py', 'NEW - LinkedIn API v2 watcher'),
    ]
    for row_idx, (filepath, status) in enumerate(files):
        cell1 = table2.rows[row_idx + 1].cells[0]
        cell2 = table2.rows[row_idx + 1].cells[1]
        if row_idx % 2 == 0:
            set_cell_shading(cell1, 'F8F8F8')
            set_cell_shading(cell2, 'F8F8F8')

        p = cell1.paragraphs[0]
        run = p.add_run(filepath)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

        p = cell2.paragraphs[0]
        run = p.add_run(status)
        run.bold = True
        run.font.size = Pt(9)
        if 'NEW' in status:
            run.font.color.rgb = RGBColor(0, 130, 0)
        else:
            run.font.color.rgb = RGBColor(0, 80, 160)

    doc.add_paragraph()

    # Recommended order
    add_styled_paragraph(doc, 'Recommended Order',
                         bold=True, font_size=14, color=(40, 40, 40), space_after=6)

    p = doc.add_paragraph('Start with the easiest and work your way up:')
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('1. Gmail ')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 130, 0)
    run = p.add_run('(easiest - just Google Cloud Console + run auth script)')
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run('2. LinkedIn ')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(200, 140, 0)
    run = p.add_run('(medium - create developer app + OAuth flow)')
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run('3. WhatsApp ')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(200, 0, 0)
    run = p.add_run('(hardest - Meta Business verification + webhook setup)')
    run.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('_' * 60)
    run.font.color.rgb = RGBColor(0, 120, 200)

    add_styled_paragraph(doc, 'Generated for AI Employee Hackathon 0 - Silver Tier',
                         italic=True, font_size=10, color=(100, 100, 100),
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, 'If you get stuck on any step, ask for help!',
                         italic=True, font_size=10, color=(100, 100, 100),
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Save
    output_path = r'D:\syeda Gulzar Bano\AI_Employee_Vault_\Watcher_Auth_Fix_Guide.docx'
    doc.save(output_path)
    print(f"Word document saved to: {output_path}")
    return output_path


if __name__ == '__main__':
    create_docx()
